"""
File Processor for vLLM Gradio WebUI

Handles file upload, processing, and content extraction for various file types
including PDFs, Word documents, images, and text files.
"""

import logging
import os
import time
import hashlib
import mimetypes
from typing import List, Dict, Optional, Any, Union, BinaryIO
from dataclasses import dataclass
from pathlib import Path
import tempfile

logger = logging.getLogger(__name__)

@dataclass
class FileInfo:
    """File information structure"""
    filename: str
    filepath: str
    file_type: str
    mime_type: str
    size: int
    content: str
    metadata: Dict[str, Any]
    processing_time: float
    success: bool
    error_message: Optional[str] = None

@dataclass
class ProcessingConfig:
    """Configuration for file processing"""
    max_file_size: int = 50 * 1024 * 1024  # 50MB
    allowed_extensions: List[str] = None
    extract_images: bool = True
    ocr_enabled: bool = False
    temp_dir: str = None

class FileProcessor:
    """Main file processor for various document types"""
    
    def __init__(self, config: ProcessingConfig = None):
        self.config = config or ProcessingConfig()
        
        # Set default allowed extensions
        if self.config.allowed_extensions is None:
            self.config.allowed_extensions = [
                '.txt', '.md', '.pdf', '.docx', '.doc', '.odt', '.rtf',
                '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff',
                '.csv', '.json', '.xml', '.html', '.htm'
            ]
        
        # Set temp directory
        if self.config.temp_dir is None:
            self.config.temp_dir = tempfile.gettempdir()
        
        # Ensure temp directory exists
        os.makedirs(self.config.temp_dir, exist_ok=True)
        
        logger.info("File processor initialized")
    
    async def process_file(self, file_path: str) -> FileInfo:
        """Process a single file"""
        start_time = time.time()
        
        try:
            # Validate file
            if not os.path.exists(file_path):
                return FileInfo(
                    filename=os.path.basename(file_path),
                    filepath=file_path,
                    file_type="unknown",
                    mime_type="",
                    size=0,
                    content="",
                    metadata={},
                    processing_time=0,
                    success=False,
                    error_message="File not found"
                )
            
            # Get file info
            file_stat = os.stat(file_path)
            filename = os.path.basename(file_path)
            file_size = file_stat.st_size
            
            # Check file size
            if file_size > self.config.max_file_size:
                return FileInfo(
                    filename=filename,
                    filepath=file_path,
                    file_type="unknown",
                    mime_type="",
                    size=file_size,
                    content="",
                    metadata={},
                    processing_time=time.time() - start_time,
                    success=False,
                    error_message=f"File too large: {file_size} bytes"
                )
            
            # Get file extension and MIME type
            file_ext = Path(file_path).suffix.lower()
            mime_type, _ = mimetypes.guess_type(file_path)
            
            # Check allowed extensions
            if file_ext not in self.config.allowed_extensions:
                return FileInfo(
                    filename=filename,
                    filepath=file_path,
                    file_type=file_ext,
                    mime_type=mime_type or "",
                    size=file_size,
                    content="",
                    metadata={},
                    processing_time=time.time() - start_time,
                    success=False,
                    error_message=f"File type not allowed: {file_ext}"
                )
            
            # Process based on file type
            content, metadata = await self._process_by_type(file_path, file_ext, mime_type)
            
            processing_time = time.time() - start_time
            
            return FileInfo(
                filename=filename,
                filepath=file_path,
                file_type=file_ext,
                mime_type=mime_type or "",
                size=file_size,
                content=content,
                metadata=metadata,
                processing_time=processing_time,
                success=True
            )
            
        except Exception as e:
            processing_time = time.time() - start_time
            logger.error(f"Failed to process file {file_path}: {e}", exc_info=True)
            
            return FileInfo(
                filename=os.path.basename(file_path),
                filepath=file_path,
                file_type="unknown",
                mime_type="",
                size=0,
                content="",
                metadata={},
                processing_time=processing_time,
                success=False,
                error_message=str(e)
            )
    
    async def process_files(self, file_paths: List[str]) -> List[FileInfo]:
        """Process multiple files"""
        import asyncio
        
        tasks = [self.process_file(file_path) for file_path in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Handle exceptions
        processed_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                processed_results.append(FileInfo(
                    filename=os.path.basename(file_paths[i]),
                    filepath=file_paths[i],
                    file_type="unknown",
                    mime_type="",
                    size=0,
                    content="",
                    metadata={},
                    processing_time=0,
                    success=False,
                    error_message=str(result)
                ))
            else:
                processed_results.append(result)
        
        logger.info(f"Processed {len(file_paths)} files, {sum(1 for r in processed_results if r.success)} successful")
        return processed_results
    
    async def _process_by_type(self, file_path: str, file_ext: str, 
                              mime_type: str) -> tuple[str, Dict[str, Any]]:
        """Process file based on its type"""
        content = ""
        metadata = {}
        
        try:
            if file_ext in ['.txt', '.md']:
                content, metadata = await self._process_text_file(file_path)
            elif file_ext == '.pdf':
                content, metadata = await self._process_pdf_file(file_path)
            elif file_ext in ['.docx', '.doc', '.odt', '.rtf']:
                content, metadata = await self._process_document_file(file_path)
            elif file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
                content, metadata = await self._process_image_file(file_path)
            elif file_ext == '.csv':
                content, metadata = await self._process_csv_file(file_path)
            elif file_ext == '.json':
                content, metadata = await self._process_json_file(file_path)
            elif file_ext in ['.xml', '.html', '.htm']:
                content, metadata = await self._process_markup_file(file_path)
            else:
                # Try to process as text
                content, metadata = await self._process_text_file(file_path)
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Failed to process {file_ext} file: {e}")
            raise
    
    async def _process_text_file(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            metadata = {
                'type': 'text',
                'encoding': 'utf-8',
                'lines': len(content.splitlines()),
                'characters': len(content)
            }
            
            return content, metadata
            
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    
                    metadata = {
                        'type': 'text',
                        'encoding': encoding,
                        'lines': len(content.splitlines()),
                        'characters': len(content)
                    }
                    
                    return content, metadata
                except:
                    continue
            
            raise ValueError("Unable to decode text file")
    
    async def _process_pdf_file(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process PDF files with enhanced image and OCR support"""
        try:
            import PyPDF2
            
            content_parts = []
            metadata = {'type': 'pdf', 'pages': 0, 'extraction_method': []}
            
            # First try standard text extraction
            try:
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    metadata['pages'] = len(pdf_reader.pages)
                    
                    # Extract text from each page
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text.strip():
                                content_parts.append(f"[Page {page_num + 1}]\n{page_text}")
                                metadata['extraction_method'].append(f"page_{page_num + 1}_text")
                            else:
                                # If no text found, mark for OCR
                                content_parts.append(f"[Page {page_num + 1}]\n[No text extracted - will try OCR]")
                        except Exception as e:
                            logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                            content_parts.append(f"[Page {page_num + 1}]\n[Text extraction failed - will try OCR]")
                    
                    # Extract metadata
                    if pdf_reader.metadata:
                        metadata.update({
                            'title': pdf_reader.metadata.get('/Title', ''),
                            'author': pdf_reader.metadata.get('/Author', ''),
                            'subject': pdf_reader.metadata.get('/Subject', ''),
                            'creator': pdf_reader.metadata.get('/Creator', ''),
                            'producer': pdf_reader.metadata.get('/Producer', ''),
                            'creation_date': str(pdf_reader.metadata.get('/CreationDate', '')),
                            'modification_date': str(pdf_reader.metadata.get('/ModDate', ''))
                        })
                        
            except Exception as e:
                logger.warning(f"PyPDF2 extraction failed: {e}, trying OCR fallback")
                content_parts = []
            
            # If OCR is enabled and we have pages with no text or extraction failed, try OCR
            if self.config.ocr_enabled and (not content_parts or any("[No text extracted" in part or "[Text extraction failed" in part for part in content_parts)):
                try:
                    from pdf2image import convert_from_path
                    import pytesseract
                    
                    logger.info(f"Converting PDF to images for OCR: {file_path}")
                    
                    # Convert PDF to images
                    images = convert_from_path(file_path, dpi=200)
                    metadata['pages'] = len(images)
                    
                    # Clear previous content if we're doing full OCR
                    if not content_parts or all("[No text extracted" in part or "[Text extraction failed" in part for part in content_parts):
                        content_parts = []
                    
                    # OCR each page
                    for page_num, image in enumerate(images):
                        try:
                            # Perform OCR
                            ocr_text = pytesseract.image_to_string(image, lang='eng')
                            if ocr_text.strip():
                                # Replace failed extraction with OCR result
                                page_content = f"[Page {page_num + 1} - OCR]\n{ocr_text}"
                                if page_num < len(content_parts):
                                    content_parts[page_num] = page_content
                                else:
                                    content_parts.append(page_content)
                                metadata['extraction_method'].append(f"page_{page_num + 1}_ocr")
                            else:
                                if page_num >= len(content_parts):
                                    content_parts.append(f"[Page {page_num + 1}]\n[No text found even with OCR]")
                        except Exception as e:
                            logger.warning(f"OCR failed for page {page_num + 1}: {e}")
                            if page_num >= len(content_parts):
                                content_parts.append(f"[Page {page_num + 1}]\n[OCR failed: {str(e)}]")
                    
                    metadata['ocr_used'] = True
                    
                except ImportError as e:
                    logger.warning(f"OCR dependencies not available: {e}")
                    metadata['ocr_error'] = "Missing dependencies: pdf2image or pytesseract"
                except Exception as e:
                    logger.warning(f"OCR processing failed: {e}")
                    metadata['ocr_error'] = str(e)
            
            content = '\n\n'.join(content_parts) if content_parts else "[No content could be extracted from PDF]"
            return content, metadata
            
        except ImportError:
            logger.error("PyPDF2 not installed. Please install with: pip install PyPDF2")
            raise
        except Exception as e:
            logger.error(f"Failed to process PDF: {e}")
            raise
    
    async def _process_document_file(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process Word documents and other document formats"""
        try:
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext == '.docx':
                return await self._process_docx_file(file_path)
            elif file_ext in ['.doc', '.odt', '.rtf']:
                # For these formats, we'd need additional libraries
                # For now, return empty content with metadata
                metadata = {
                    'type': 'document',
                    'format': file_ext,
                    'note': f'{file_ext} format not fully supported'
                }
                return "", metadata
            else:
                raise ValueError(f"Unsupported document format: {file_ext}")
                
        except Exception as e:
            logger.error(f"Failed to process document: {e}")
            raise
    
    async def _process_docx_file(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process DOCX files"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Extract text from paragraphs
            content_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content_parts.append(' | '.join(row_text))
            
            content = '\n\n'.join(content_parts)
            
            # Extract metadata
            metadata = {
                'type': 'docx',
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
            
            # Extract document properties
            if doc.core_properties:
                metadata.update({
                    'title': doc.core_properties.title or '',
                    'author': doc.core_properties.author or '',
                    'subject': doc.core_properties.subject or '',
                    'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                    'modified': str(doc.core_properties.modified) if doc.core_properties.modified else ''
                })
            
            return content, metadata
            
        except ImportError:
            logger.error("python-docx not installed. Please install with: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Failed to process DOCX: {e}")
            raise
    
    async def _process_image_file(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process image files"""
        try:
            from PIL import Image
            
            with Image.open(file_path) as img:
                metadata = {
                    'type': 'image',
                    'format': img.format,
                    'mode': img.mode,
                    'size': img.size,
                    'width': img.width,
                    'height': img.height
                }
                
                # Extract EXIF data if available
                if hasattr(img, '_getexif') and img._getexif():
                    exif_data = img._getexif()
                    if exif_data:
                        metadata['exif'] = {str(k): str(v) for k, v in exif_data.items()}
            
            # OCR processing if enabled
            content = ""
            if self.config.ocr_enabled:
                content = await self._perform_ocr(file_path)
                metadata['ocr_performed'] = True
            else:
                content = f"Image file: {os.path.basename(file_path)}"
                metadata['ocr_performed'] = False
            
            return content, metadata
            
        except ImportError:
            logger.error("Pillow not installed. Please install with: pip install Pillow")
            raise
        except Exception as e:
            logger.error(f"Failed to process image: {e}")
            raise
    
    async def _perform_ocr(self, file_path: str) -> str:
        """Perform OCR on image file"""
        try:
            import pytesseract
            from PIL import Image
            
            with Image.open(file_path) as img:
                text = pytesseract.image_to_string(img)
                return text.strip()
                
        except ImportError:
            logger.warning("pytesseract not installed. OCR disabled.")
            return f"Image file: {os.path.basename(file_path)} (OCR not available)"
        except Exception as e:
            logger.warning(f"OCR failed: {e}")
            return f"Image file: {os.path.basename(file_path)} (OCR failed)"
    
    async def _process_csv_file(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process CSV files"""
        try:
            import csv
            
            content_parts = []
            row_count = 0
            
            with open(file_path, 'r', encoding='utf-8') as f:
                # Try to detect delimiter
                sample = f.read(1024)
                f.seek(0)
                
                sniffer = csv.Sniffer()
                delimiter = sniffer.sniff(sample).delimiter
                
                reader = csv.reader(f, delimiter=delimiter)
                
                for i, row in enumerate(reader):
                    if i == 0:
                        # Header row
                        content_parts.append(f"Headers: {', '.join(row)}")
                    else:
                        # Data rows (limit to first 100 rows for content)
                        if i <= 100:
                            content_parts.append(' | '.join(row))
                    row_count = i + 1
            
            content = '\n'.join(content_parts)
            
            metadata = {
                'type': 'csv',
                'rows': row_count,
                'delimiter': delimiter
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Failed to process CSV: {e}")
            raise
    
    async def _process_json_file(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process JSON files"""
        try:
            import json
            
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert JSON to readable text
            content = json.dumps(data, indent=2, ensure_ascii=False)
            
            metadata = {
                'type': 'json',
                'structure': type(data).__name__
            }
            
            if isinstance(data, dict):
                metadata['keys'] = list(data.keys())[:10]  # First 10 keys
            elif isinstance(data, list):
                metadata['items'] = len(data)
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Failed to process JSON: {e}")
            raise
    
    async def _process_markup_file(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process XML/HTML files"""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Parse with BeautifulSoup
            soup = BeautifulSoup(content, 'html.parser')
            
            # Extract text content
            text_content = soup.get_text()
            
            # Clean up whitespace
            text_content = '\n'.join(line.strip() for line in text_content.splitlines() if line.strip())
            
            metadata = {
                'type': 'markup',
                'format': Path(file_path).suffix.lower(),
                'original_length': len(content),
                'text_length': len(text_content)
            }
            
            # Extract title if HTML
            if soup.title:
                metadata['title'] = soup.title.string
            
            return text_content, metadata
            
        except ImportError:
            logger.error("BeautifulSoup not installed. Please install with: pip install beautifulsoup4")
            # Fallback to plain text processing
            return await self._process_text_file(file_path)
        except Exception as e:
            logger.error(f"Failed to process markup: {e}")
            raise
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions"""
        return self.config.allowed_extensions.copy()
    
    def is_supported_file(self, filename: str) -> bool:
        """Check if file is supported"""
        file_ext = Path(filename).suffix.lower()
        return file_ext in self.config.allowed_extensions
    
    def get_file_type_info(self, filename: str) -> Dict[str, Any]:
        """Get information about file type"""
        file_ext = Path(filename).suffix.lower()
        mime_type, _ = mimetypes.guess_type(filename)
        
        type_categories = {
            'text': ['.txt', '.md'],
            'document': ['.pdf', '.docx', '.doc', '.odt', '.rtf'],
            'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff'],
            'data': ['.csv', '.json', '.xml'],
            'web': ['.html', '.htm']
        }
        
        category = 'unknown'
        for cat, extensions in type_categories.items():
            if file_ext in extensions:
                category = cat
                break
        
        return {
            'extension': file_ext,
            'mime_type': mime_type,
            'category': category,
            'supported': self.is_supported_file(filename)
        }

